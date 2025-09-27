import json
import random
import datetime
from docx import Document
from docx.shared import Inches
import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import argparse
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Centralized constants
from variables import CFG_FOLDER, OUTPUT_FOLDER

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def main():

    parser = argparse.ArgumentParser()
    parser.add_argument(
        "-i", "--input", type=str, required=False, help="Please enter a input json in cfg folder"
    )
    parser.add_argument(
        "--gui", action="store_true", help="Enable GUI mode"
    )
    args = parser.parse_args()

    if args.gui:
        run_gui()
    else:
        if args.input:
            _main(args.input)
        else:
            print("Please provide an input file with -i or use --gui for GUI mode")


def _main(read_from_filename="1A-P12.json"):

    if read_from_filename.endswith(".json"):
        read_from_filename = read_from_filename
    else:
        read_from_filename = read_from_filename + ".json"

    base_filename = read_from_filename.split(".")[0]

    file_path = os.path.join(CFG_FOLDER, read_from_filename)
    with open(file_path, "r", encoding="utf8") as f:
        data = json.load(f)
    lst_explain = data.get("explain", [])
    lst_statement = data["statement"]

    # 將 lst_explain 與 lst_statement 的打亂

    random.seed(datetime.datetime.now().timestamp())

    random.shuffle(lst_explain)
    random.shuffle(lst_statement)
    print(lst_explain)
    print(lst_statement)

    ans_doc = Document()
    document = Document()

    headings = ["1st", "2nd"]
    # document.add_heading('TEST', 0)
    heading_count = 0
    if len(lst_explain) > 0:
        document.add_heading(headings[heading_count], 1)
        ans_doc.add_heading(headings[heading_count], 1)
        heading_count += 1
        i = 0
        for e, w in lst_explain:
            i += 1
            document.add_paragraph(f"{i} ________________: {e}")
        i = 0
        for e, w in lst_explain:
            i+=1
            ans_doc.add_paragraph(f"{i} {w} : {e}")

    if len(lst_statement) > 0:
        document.add_heading(headings[heading_count], 1)
        ans_doc.add_heading(headings[heading_count], 1)
        heading_count += 1
        i = 0
        for s, w in lst_statement:
            i += 1
            if w not in s:
                raise ValueError(f"Word {w} not in statement {s}")
            else:
                document.add_paragraph(f'{i} {s.replace(w, "__________________")}')
        i = 0
        for s, w in lst_statement:
            i += 1
            ans_doc.add_paragraph(f"{i} {w} : {s}")

    def change_font_size(doc, font_size):
        # Change the font size of all paragraphs in the document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                font_name = "Comic Sans MS"
                run.font.name = font_name
        return doc

    change_font_size(document, 12)

    # Set document margins to 2 cm
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1 / 2.54)
        section.bottom_margin = Inches(1 / 2.54)
        section.left_margin = Inches(1 / 2.54)
        section.right_margin = Inches(1 / 2.54)

    k = 0

    # 將 doc
    while True:
        filename = f"{base_filename}_test-{k}.docx"
        ans_filename = f"{base_filename}_test-{k}-ans.docx"
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        ans_filepath = os.path.join(OUTPUT_FOLDER, ans_filename)

        if not os.path.exists(filepath):
            break
        k += 1
    # 將頁首加上 hello world
    header = document.sections[0].header
    header.paragraphs[0].text = filename

    header = ans_doc.sections[0].header
    header.paragraphs[0].text = ans_filename

    ans_doc.save(ans_filepath)
    document.save(filepath)

    return filepath, ans_filepath


def run_gui():
    def generate_file(print_file=False):
        selected_folder = folder_var.get()
        if not selected_folder:
            messagebox.showerror("Error", "Please select a file")
            return
        filepath, ans_filepath = _main(selected_folder)
        if print_file:
            os.startfile(filepath, "print")

    root = tk.Tk()
    root.title("Word Test Paper Generator")
    root.geometry("400x300")
    
    # Windows 工作列圖示設定 - 簡化但有效的方法
    try:
        # 設定應用程式 ID（必須在圖示設定前）
        import ctypes
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("WordTestPaperGenerator.1.0")
    except:
        pass
    
    # 設定圖示
    png_path = os.path.abspath(os.path.join("icons", "score.png"))
    ico_path = os.path.abspath(os.path.join("icons", "score.ico"))
    
    # 如果有 PNG 檔案，強制重新創建 ICO 檔案
    if PIL_AVAILABLE and os.path.exists(png_path):
        try:
            img = Image.open(png_path)
            if img.mode != 'RGBA':
                img = img.convert('RGBA')
            
            # 刪除舊的 ICO 檔案並重新創建
            if os.path.exists(ico_path):
                os.remove(ico_path)
            
            # 創建標準尺寸的 ICO 檔案
            sizes = [(16, 16), (32, 32), (48, 48), (256, 256)]
            img.save(ico_path, format='ICO', sizes=sizes)
            print("重新創建了 ICO 檔案")
        except Exception as e:
            print(f"創建 ICO 失敗: {e}")
    
    # 設定視窗圖示
    icon_set = False
    if os.path.exists(ico_path):
        try:
            root.iconbitmap(ico_path)
            icon_set = True
            print("設定了 ICO 圖示")
        except Exception as e:
            print(f"ICO 設定失敗: {e}")
    
    # 同時設定 PhotoImage 圖示（雙重保險）
    if os.path.exists(png_path):
        try:
            icon_img = tk.PhotoImage(file=png_path)
            root.iconphoto(True, icon_img)
            root.icon_ref = icon_img  # 保持引用
            print("設定了 PNG 圖示")
        except Exception as e:
            print(f"PNG 設定失敗: {e}")
    
    # 強制刷新視窗
    root.update()
    
    print(f"圖示檔案路徑: ICO={ico_path}, PNG={png_path}")
    print(f"ICO 存在: {os.path.exists(ico_path)}, PNG 存在: {os.path.exists(png_path)}")
    
    root.option_add("*Font", "Arial 16")

    folder_var = tk.StringVar()
    folder_label = ttk.Label(root, text="Select file:")
    folder_label.pack(pady=5)
    
    folder_dropdown = ttk.Combobox(root, textvariable=folder_var)
    folder_dropdown.pack(pady=5)
    
    folders = [f for f in os.listdir(CFG_FOLDER) if f.endswith(".json")]
    folder_dropdown["values"] = folders

    # 建立一個樣式
    style = ttk.Style()
    style.configure("TButton", font=("Arial", 14))


    generate_button = ttk.Button(root, text="Generate File", command=lambda: generate_file(print_file=False), style="TButton")   
    generate_button.pack(side="left", padx=20)
    
    generate_print_button = ttk.Button(root, text="Generate and Print", command=lambda: generate_file(print_file=True),style="TButton") 
    generate_print_button.pack(side="right", padx=20)
    

    root.mainloop()


if __name__ == "__main__":
    main()
