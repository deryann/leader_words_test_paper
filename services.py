"""
Service classes implementing the core business logic.
Each class has a single responsibility following SOLID principles.
"""
import json
import os
import random
import datetime
from typing import List
from docx import Document
from docx.shared import Inches, Pt

from interfaces import (
    ConfigLoaderInterface, DocumentGeneratorInterface, 
    FileManagerInterface, DataShufflerInterface
)
from models import TestData, TestItem, TestPaperConfig, GeneratedFiles
from exceptions import ConfigurationError, ValidationError, DocumentGenerationError
from variables import CFG_FOLDER, OUTPUT_FOLDER


class ConfigLoader(ConfigLoaderInterface):
    """Handles loading and parsing of test configuration files."""
    
    def __init__(self, config_folder: str = CFG_FOLDER):
        self.config_folder = config_folder
    
    def load_config(self, filename: str) -> TestData:
        """Load test data from JSON configuration file."""
        if not filename.endswith(".json"):
            filename += ".json"
        
        file_path = os.path.join(self.config_folder, filename)
        
        if not os.path.exists(file_path):
            raise ConfigurationError(f"Configuration file not found: {file_path}")
        
        try:
            with open(file_path, "r", encoding="utf8") as f:
                data = json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            raise ConfigurationError(f"Error reading configuration file: {e}")
        
        # Parse explain items
        explain_items = []
        for item in data.get("explain", []):
            if isinstance(item, (list, tuple)) and len(item) == 2:
                explain_items.append(TestItem(text=item[0], word=item[1]))
            else:
                raise ConfigurationError(f"Invalid explain item format: {item}")
        
        # Parse statement items
        statement_items = []
        for item in data.get("statement", []):
            if isinstance(item, (list, tuple)) and len(item) == 2:
                statement_items.append(TestItem(text=item[0], word=item[1]))
            else:
                raise ConfigurationError(f"Invalid statement item format: {item}")
        
        test_data = TestData(explain_items=explain_items, statement_items=statement_items)
        
        if not test_data.validate():
            raise ValidationError("Test data validation failed")
        
        return test_data
    
    def get_available_files(self) -> List[str]:
        """Get list of available JSON configuration files."""
        if not os.path.exists(self.config_folder):
            return []
        
        return [f for f in os.listdir(self.config_folder) if f.endswith(".json")]


class FileManager(FileManagerInterface):
    """Manages file operations and path handling."""
    
    def __init__(self, output_folder: str = OUTPUT_FOLDER):
        self.output_folder = output_folder
        self.ensure_output_directory()
    
    def ensure_output_directory(self) -> None:
        """Ensure output directory exists."""
        os.makedirs(self.output_folder, exist_ok=True)
    
    def get_unique_filename(self, base_filename: str, extension: str = ".docx") -> str:
        """Generate a unique filename to avoid conflicts."""
        counter = 0
        while True:
            if counter == 0:
                filename = f"{base_filename}{extension}"
            else:
                filename = f"{base_filename}-{counter}{extension}"
            
            filepath = os.path.join(self.output_folder, filename)
            if not os.path.exists(filepath):
                return filename
            counter += 1


class DataShuffler(DataShufflerInterface):
    """Handles randomization of test data."""
    
    def shuffle_data(self, test_data: TestData) -> TestData:
        """Shuffle test data items randomly using current timestamp as seed."""
        random.seed(datetime.datetime.now().timestamp())
        
        # Create copies to avoid modifying original data
        shuffled_explain = test_data.explain_items.copy()
        shuffled_statement = test_data.statement_items.copy()
        
        random.shuffle(shuffled_explain)
        random.shuffle(shuffled_statement)
        
        return TestData(
            explain_items=shuffled_explain,
            statement_items=shuffled_statement
        )


class DocumentGenerator(DocumentGeneratorInterface):
    """Generates Word documents for test papers and answer sheets."""
    
    def __init__(self, file_manager: FileManagerInterface):
        self.file_manager = file_manager
    
    def generate_test_paper(self, test_data: TestData, config: TestPaperConfig) -> GeneratedFiles:
        """Generate test paper and answer sheet documents."""
        try:
            base_filename = config.input_filename.split(".")[0]
            
            # Generate unique filenames  
            test_counter = 0
            ans_counter = 0
            
            # Find unique test filename
            while True:
                if test_counter == 0:
                    test_filename = f"{base_filename}_test.docx"
                else:
                    test_filename = f"{base_filename}_test-{test_counter}.docx"
                
                test_filepath = os.path.join(self.file_manager.output_folder, test_filename)
                if not os.path.exists(test_filepath):
                    break
                test_counter += 1
            
            # Find unique answer filename
            while True:
                if ans_counter == 0:
                    ans_filename = f"{base_filename}_test-{test_counter}-ans.docx"
                else:
                    ans_filename = f"{base_filename}_test-{test_counter}-ans-{ans_counter}.docx"
                
                ans_filepath = os.path.join(self.file_manager.output_folder, ans_filename)
                if not os.path.exists(ans_filepath):
                    break
                ans_counter += 1
            
            # Create documents
            test_doc = Document()
            ans_doc = Document()
            
            # Generate content
            self._generate_document_content(test_doc, ans_doc, test_data, config)
            
            # Apply formatting
            self._apply_document_formatting(test_doc, config)
            self._apply_document_formatting(ans_doc, config)
            
            # Set headers
            self._set_document_header(test_doc, test_filename)
            self._set_document_header(ans_doc, ans_filename)
            
            # Save documents
            test_doc.save(test_filepath)
            ans_doc.save(ans_filepath)
            
            return GeneratedFiles(
                test_file_path=test_filepath,
                answer_file_path=ans_filepath,
                base_filename=base_filename
            )
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to generate documents: {e}")
    
    def _generate_document_content(self, test_doc: Document, ans_doc: Document, 
                                 test_data: TestData, config: TestPaperConfig) -> None:
        """Generate content for both test and answer documents."""
        heading_count = 0
        
        # Generate explain section
        if test_data.explain_items:
            test_doc.add_heading(config.headings[heading_count], 1)
            ans_doc.add_heading(config.headings[heading_count], 1)
            heading_count += 1
            
            for i, item in enumerate(test_data.explain_items, 1):
                test_doc.add_paragraph(f"{i} ________________: {item.text}")
                ans_doc.add_paragraph(f"{i} {item.word} : {item.text}")
        
        # Generate statement section
        if test_data.statement_items:
            test_doc.add_heading(config.headings[heading_count], 1)
            ans_doc.add_heading(config.headings[heading_count], 1)
            
            for i, item in enumerate(test_data.statement_items, 1):
                test_text = item.text.replace(item.word, "__________________")
                test_doc.add_paragraph(f"{i} {test_text}")
                ans_doc.add_paragraph(f"{i} {item.word} : {item.text}")
    
    def _apply_document_formatting(self, doc: Document, config: TestPaperConfig) -> None:
        """Apply formatting to document."""
        # Set font for all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(config.font_size)
                run.font.name = config.font_name
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(config.margin_inches)
            section.bottom_margin = Inches(config.margin_inches)
            section.left_margin = Inches(config.margin_inches)
            section.right_margin = Inches(config.margin_inches)
    
    def _set_document_header(self, doc: Document, filename: str) -> None:
        """Set document header with filename."""
        header = doc.sections[0].header
        header.paragraphs[0].text = filename